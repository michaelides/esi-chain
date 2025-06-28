import chainlit as cl
import os
import json
import re
import uuid
# import extra_streamlit_components as esc # Removed Streamlit specific component
from typing import List, Dict, Any, Optional, AsyncGenerator
from llama_index.core.llms import ChatMessage, MessageRole
from llama_index.core import Settings
# import stui # Removed Streamlit UI import
from agent import create_orchestrator_agent, generate_suggested_prompts, SUGGESTED_PROMPT_COUNT, DEFAULT_PROMPTS, initialize_settings as initialize_agent_settings, generate_llm_greeting
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
# import io # Already imported BytesIO
from llama_index.core.tools import FunctionTool

from huggingface_hub import HfFileSystem
from config import HF_TOKEN
# import os # Already imported

# Import ui.py to make Chainlit decorators discoverable
import ui # IMPORTANT

fs = HfFileSystem()
load_dotenv()

PROJECT_ROOT = os.path.abspath(os.path.dirname(__file__))

# cookies = esc.CookieManager(key="esi_cookie_manager") # Removed Streamlit specific cookie manager
# Chainlit handles sessions differently. User identity might be managed by Chainlit or custom auth.

SIMPLE_STORE_PATH_RELATIVE = os.getenv("SIMPLE_STORE_PATH", "ragdb/simple_vector_store")
DB_PATH = os.path.join(PROJECT_ROOT, SIMPLE_STORE_PATH_RELATIVE)
# AGENT_SESSION_KEY = "esi_orchestrator_agent" # Will be part of AppState
DOWNLOAD_MARKER = "---DOWNLOAD_FILE---" # Still used for formatting messages
RAG_SOURCE_MARKER_PREFIX = "---RAG_SOURCE---" # Still used for formatting messages

from tools import UI_ACCESSIBLE_WORKSPACE
from config import HF_USER_MEMORIES_DATASET_ID

MAX_CHAT_HISTORY_MESSAGES = 15

# Global flag to ensure LLM settings are initialized once per application lifecycle
_global_llm_settings_initialized = False
_global_llm_initialization_error: Optional[str] = None

class AppState:
    def __init__(self, user_session): # Removed cl.UserSession type hint
        self.user_session = user_session # Store the Chainlit user session
        self.user_id: Optional[str] = None
        self.agent: Optional[Any] = None
        self.chat_metadata: Dict[str, str] = {}
        self.all_chat_messages: Dict[str, List[Dict[str, Any]]] = {}
        self.current_chat_id: Optional[str] = None
        self.messages: List[Dict[str, Any]] = [] # Current chat messages
        self.suggested_prompts: List[str] = DEFAULT_PROMPTS.copy()

        self.uploaded_documents: Dict[str, str] = {} # filename: content
        self.uploaded_dataframes: Dict[str, Any] = {} # filename: dataframe (consider storing path/summary)

        # LLM Settings - to be loaded/managed
        self.llm_temperature: float = 0.7
        self.llm_verbosity: int = 3
        self.search_results_count: int = 5
        self.long_term_memory_enabled: bool = True # Default, will be loaded from user preference

        self.chat_modified: bool = False # Tracks if current chat needs saving
        self.initial_greeting_text: str = "" # Will be generated

        # Ensure global LLM settings are set up
        global _global_llm_settings_initialized, _global_llm_initialization_error
        if not _global_llm_settings_initialized and not _global_llm_initialization_error:
            try:
                initialize_agent_settings()
                _global_llm_settings_initialized = True
                print("Global LLM settings initialized successfully.")
            except Exception as e:
                _global_llm_initialization_error = f"Fatal Error: Could not initialize LLM settings. {e}"
                print(_global_llm_initialization_error)

        if _global_llm_initialization_error:
            # If there's a global error, we might want to prevent agent setup or signal error to UI
            print(f"AppState initialized with global LLM error: {_global_llm_initialization_error}")


    async def initialize_session(self):
        """Async initialization steps that depend on external factors or need async ops."""
        print("AppState: Initializing session...")

        if _global_llm_initialization_error:
            await cl.ErrorMessage(content=_global_llm_initialization_error, author="System").send()
            # Potentially block further operations if LLM is critical

        # User ID and Long-Term Memory Preference
        # Chainlit's cl.User might provide a stable ID if auth is used.
        # For simplicity, we'll mimic cookie behavior using user_session for preference.
        # A more robust solution would use Chainlit's user identification if available.

        pref_from_session = self.user_session.get("long_term_memory_pref")
        if pref_from_session is not None:
            self.long_term_memory_enabled = bool(pref_from_session)
        else:
            self.user_session.set("long_term_memory_pref", self.long_term_memory_enabled) # Save default

        # User ID: Chainlit provides `cl.user_session.get("id")` which is unique per session.
        # For persistent user ID across sessions (if LTM is on), we'd need custom handling
        # or rely on Chainlit's auth features if enabled.
        # For now, let's use a session-based ID and manage persistence via HF if LTM is on.
        # This part needs careful thought for production to correctly map users to their HF data.

        # Simplified user ID logic for now:
        # If LTM is on, try to get a persistent ID (e.g., from Chainlit user object if auth is on, or a custom cookie)
        # If LTM is off, or no persistent ID, use a session-specific ID.
        self.user_id = self.user_session.get("id") # This is a session ID.
        # TODO: Implement robust user ID management for LTM.
        # For now, user_id will be the Chainlit session ID. This means LTM won't truly persist across browser closes
        # without further work on user identification.

        if self.long_term_memory_enabled and self.user_id: # Check user_id exists
            print(f"Loading user data for user {self.user_id} from Hugging Face...")
            user_data = self._load_user_data_from_hf(self.user_id)
            self.chat_metadata = user_data.get("metadata", {})
            self.all_chat_messages = user_data.get("messages", {})
            print(f"Loaded {len(self.chat_metadata)} chats for user {self.user_id}.")
        else:
            print(f"Long-term memory disabled or no user_id. No historical data loaded for user_id {self.user_id}.")

        # Initialize LLM settings from user_session or defaults
        self.llm_temperature = self.user_session.get("llm_temperature", self.llm_temperature)
        self.llm_verbosity = self.user_session.get("llm_verbosity", self.llm_verbosity)
        self.search_results_count = self.user_session.get("search_results_count", self.search_results_count)

        self.initial_greeting_text = generate_llm_greeting() # Generate once

        # Setup agent
        await self._setup_agent_async() # Make it async if agent setup can be slow

        # Active Chat and Initial Greeting Logic
        if self.long_term_memory_enabled:
            if self.chat_metadata: # If there are saved chats
                # Try to load last active chat ID if stored, or default to first
                last_chat_id = self.user_session.get("current_chat_id")
                if last_chat_id and last_chat_id in self.chat_metadata:
                    self.current_chat_id = last_chat_id
                else:
                    self.current_chat_id = next(iter(self.chat_metadata)) # First chat

                self.messages = self.all_chat_messages.get(self.current_chat_id, [])
                self.chat_modified = True # Assume loaded chat is "modified" in terms of session state
            else: # No saved chats, start new
                await self.create_new_chat(is_initial_greeting=True)
        else: # LTM disabled, always start new temporary chat
            await self.create_new_chat(is_initial_greeting=True)

        self.suggested_prompts = await self.generate_suggested_prompts_async(self.messages)
        print("AppState session initialization complete.")

    async def _setup_agent_async(self):
        print("AppState: Initializing AI agent...")
        if _global_llm_initialization_error:
            print(f"Skipping agent setup due to LLM init error: {_global_llm_initialization_error}")
            self.agent = None
            return

        try:
            uploaded_doc_reader_tool = FunctionTool.from_defaults(
                fn=self._read_uploaded_document_tool_fn, # Method of AppState
                name="read_uploaded_document",
                description="Reads the full text content of a document previously uploaded by the user. Input is the exact filename (e.g., 'my_dissertation.pdf'). Use this to answer questions about the content of uploaded documents."
            )
            dataframe_analyzer_tool = FunctionTool.from_defaults(
                fn=self._analyze_dataframe_tool_fn, # Method of AppState
                name="analyze_uploaded_dataframe",
                description="Provides summary information (shape, columns, dtypes, head, describe) about a pandas DataFrame previously uploaded by the user. Input is the exact filename (e.g., 'my_data.csv'). Use this to understand the structure and basic statistics of uploaded datasets. For more complex analysis, use the 'code_interpreter' tool."
            )
            self.agent = create_orchestrator_agent(
                dynamic_tools=[uploaded_doc_reader_tool, dataframe_analyzer_tool],
                max_search_results=self.search_results_count
            )
            print("AppState: AI agent initialized successfully.")
        except Exception as e:
            error_message = f"Failed to initialize the AI agent. Please check configurations. Error: {e}"
            print(f"Error initializing AI agent: {e}")
            self.agent = None
            await cl.ErrorMessage(content=error_message, author="System").send()

    # --- Tool Functions (as methods of AppState) ---
    def _read_uploaded_document_tool_fn(self, filename: str) -> str:
        if filename not in self.uploaded_documents:
            return f"Error: Document '{filename}' not found. Available documents: {list(self.uploaded_documents.keys())}"
        return self.uploaded_documents[filename]

    def _analyze_dataframe_tool_fn(self, filename: str, head_rows: int = 5) -> str:
        if filename not in self.uploaded_dataframes:
            return f"Error: DataFrame '{filename}' not found. Available dataframes: {list(self.uploaded_dataframes.keys())}"
        df = self.uploaded_dataframes[filename]
        # ... (rest of the analysis logic from original function)
        info_str = f"DataFrame: {filename}\nShape: {df.shape}\nColumns: {', '.join(df.columns)}\n"
        info_str += f"Data Types:\n{df.dtypes.to_string()}\n"
        head_rows = max(0, min(head_rows, len(df)))
        if head_rows > 0:
            info_str += f"First {head_rows} rows:\n{df.head(head_rows).to_string()}\n"
        else:
            info_str += "No head rows requested or available.\n"
        info_str += f"Summary Statistics:\n{df.describe().to_string()}\n"
        return info_str

    # --- Hugging Face Data Management ---
    def _load_user_data_from_hf(self, user_id_to_load: str) -> Dict[str, Any]:
        hf_token = os.getenv("HF_TOKEN")
        if not hf_token:
            print("HF_TOKEN env var not set. Cannot load user data from HF.")
            return {"metadata": {}, "messages": {}}
        # ... (logic from original _load_user_data_from_hf, using user_id_to_load)
        all_chat_metadata = {}
        all_chat_messages = {}
        try:
            metadata_filename_in_repo = f"user_memories/{user_id_to_load}_metadata.json"
            messages_filename_in_repo = f"user_memories/{user_id_to_load}_messages.json"
            metadata_hf_path = f"datasets/{HF_USER_MEMORIES_DATASET_ID}/{metadata_filename_in_repo}"
            messages_hf_path = f"datasets/{HF_USER_MEMORIES_DATASET_ID}/{messages_filename_in_repo}"

            try:
                metadata_content = fs.read_text(metadata_hf_path, token=hf_token)
                all_chat_metadata = json.loads(metadata_content)
            except FileNotFoundError: # Corrected exception type
                all_chat_metadata = {} # Default to empty if not found
            except Exception as e: # Catch other potential errors during load
                print(f"Error loading metadata for user {user_id_to_load} from {metadata_hf_path}: {e}")
                all_chat_metadata = {}


            try:
                messages_content = fs.read_text(messages_hf_path, token=hf_token)
                all_chat_messages = json.loads(messages_content)
            except FileNotFoundError: # Corrected exception type
                 all_chat_messages = {} # Default to empty if not found
            except Exception as e: # Catch other potential errors during load
                print(f"Error loading messages for user {user_id_to_load} from {messages_hf_path}: {e}")
                all_chat_messages = {}

            return {"metadata": all_chat_metadata, "messages": all_chat_messages}
        except Exception as e:
            print(f"Error loading user data from HF for user {user_id_to_load}: {e}")
            return {"metadata": {}, "messages": {}}


    async def _save_chat_history_hf(self):
        if not self.long_term_memory_enabled or not self.user_id or not self.current_chat_id:
            return
        hf_token = os.getenv("HF_TOKEN")
        if not hf_token: return

        try:
            messages_filename_in_repo = f"user_memories/{self.user_id}_messages.json"
            messages_hf_path = f"datasets/{HF_USER_MEMORIES_DATASET_ID}/{messages_filename_in_repo}"

            # Load existing, update, save
            try:
                existing_content = fs.read_text(messages_hf_path, token=hf_token)
                all_user_messages = json.loads(existing_content)
            except FileNotFoundError: # Corrected exception type
                all_user_messages = {}
            except Exception: # Catch other potential errors during load
                all_user_messages = {}


            all_user_messages[self.current_chat_id] = self.messages
            with fs.open(messages_hf_path, "w", token=hf_token) as f:
                f.write(json.dumps(all_user_messages, indent=2))
            print(f"Chat history for chat {self.current_chat_id} saved to HF.")
        except Exception as e:
            print(f"Error saving chat history to HF: {e}")
            await cl.ErrorMessage(content=f"Error saving chat to cloud: {e}", author="System").send()

    async def _save_chat_metadata_hf(self):
        if not self.long_term_memory_enabled or not self.user_id:
            return
        hf_token = os.getenv("HF_TOKEN")
        if not hf_token: return

        try:
            metadata_filename_in_repo = f"user_memories/{self.user_id}_metadata.json"
            metadata_hf_path = f"datasets/{HF_USER_MEMORIES_DATASET_ID}/{metadata_filename_in_repo}"
            with fs.open(metadata_hf_path, "w", token=hf_token) as f:
                f.write(json.dumps(self.chat_metadata, indent=2))
            print(f"Chat metadata for user {self.user_id} saved to HF.")
        except Exception as e:
            print(f"Error saving chat metadata to HF: {e}")
            await cl.ErrorMessage(content=f"Error saving chat metadata to cloud: {e}", author="System").send()

    # --- Chat History Formatting ---
    def _format_chat_history_for_agent(self) -> List[ChatMessage]:
        truncated_messages = self.messages[-MAX_CHAT_HISTORY_MESSAGES:]
        history = []
        for msg in truncated_messages:
            role = MessageRole.USER if msg["role"] == "user" else MessageRole.ASSISTANT
            history.append(ChatMessage(role=role, content=msg["content"]))
        return history

    # --- Agent Interaction ---
    async def get_agent_response_stream(self, query: str) -> AsyncGenerator[str, None]:
        if not self.agent:
            yield "Agent not initialized. Cannot process query."
            return
        if _global_llm_initialization_error:
             yield f"Cannot get agent response due to LLM initialization error: {_global_llm_initialization_error}"
             return

        try:
            if Settings.llm: # Ensure Settings.llm is not None
                if hasattr(Settings.llm, 'temperature'):
                    Settings.llm.temperature = self.llm_temperature
                # else: (handle cases where temperature attribute might be missing if other models are used)
            # else: (handle cases where Settings.llm itself is None)


            modified_query = f"Verbosity Level: {self.llm_verbosity}. {query}"
            chat_history_for_agent = self._format_chat_history_for_agent() # Uses self.messages

            # LlamaIndex agent.chat is synchronous. For true async streaming with llama-index,
            # agent.astream_chat() would be used if available and the agent supports it.
            # For now, we simulate async behavior if agent.chat is blocking.
            # If create_orchestrator_agent returns an agent with a .chat method that's sync:
            response = await cl.make_async(self.agent.chat)(modified_query, chat_history=chat_history_for_agent)
            response_text = response.response if hasattr(response, 'response') else str(response)

            # Simulate streaming token by token (or word by word)
            for token in response_text.split(" "): # Simplistic word splitting
                yield token + " "
                await cl.sleep(0.02) # Small delay for streaming effect
        except Exception as e:
            error_message = f"I apologize, but I encountered an error: {str(e)}"
            print(f"Error getting agent response: {type(e).__name__} - {e}")
            yield error_message


    # --- Chat Management Methods ---
    async def create_new_chat(self, is_initial_greeting=False):
        new_chat_id = str(uuid.uuid4())
        new_chat_name = "Current Session"
        if self.long_term_memory_enabled:
            existing_idea_nums = [int(re.match(r"Idea (\d+)", name).group(1)) for name in self.chat_metadata.values() if re.match(r"Idea (\d+)", name)]
            next_idea_num = max(existing_idea_nums) + 1 if existing_idea_nums else 1
            new_chat_name = f"Idea {next_idea_num}"

        self.chat_metadata[new_chat_id] = new_chat_name
        self.messages = [{"role": "assistant", "content": self.initial_greeting_text if is_initial_greeting else "New chat started."}]
        self.all_chat_messages[new_chat_id] = self.messages
        self.current_chat_id = new_chat_id
        self.chat_modified = False # New chats are initially unsaved (metadata saved on first interaction)

        self.user_session.set("current_chat_id", self.current_chat_id) # Save to session for persistence
        if self.long_term_memory_enabled:
            await self._save_chat_metadata_hf() # Save new chat's metadata
            # History saved after first user interaction

        self.suggested_prompts = await self.generate_suggested_prompts_async(self.messages)
        print(f"Created new chat: '{new_chat_name}' (ID: {new_chat_id})")
        return new_chat_id

    async def switch_chat(self, chat_id: str):
        if not self.long_term_memory_enabled:
            await cl.Message(content="Long-term memory is disabled. Cannot switch to historical chats. Starting a new temporary session.", author="System").send()
            await self.create_new_chat(is_initial_greeting=True)
            return

        if chat_id not in self.chat_metadata:
            await cl.ErrorMessage(content=f"Error: Chat ID '{chat_id}' not found.", author="System").send()
            return

        self.current_chat_id = chat_id
        self.messages = self.all_chat_messages.get(chat_id, [])
        self.chat_modified = True # Assume loaded chat is "modified"
        self.user_session.set("current_chat_id", self.current_chat_id)
        self.suggested_prompts = await self.generate_suggested_prompts_async(self.messages)
        print(f"Switched to chat: '{self.chat_metadata.get(chat_id, 'Unknown')}' (ID: {chat_id})")
        # UI update will be handled by Chainlit rerun/message flow

    async def delete_chat(self, chat_id: str):
        if not self.long_term_memory_enabled or not self.user_id:
            # For non-LTM, "deleting" current chat means starting a new one
            if chat_id == self.current_chat_id:
                await self.create_new_chat(is_initial_greeting=True)
                await cl.Message(content="Current temporary chat cleared. Started a new one.", author="System").send()
            return

        is_current_chat = (chat_id == self.current_chat_id)

        if chat_id in self.all_chat_messages: del self.all_chat_messages[chat_id]
        if chat_id in self.chat_metadata: del self.chat_metadata[chat_id]

        await self._save_chat_metadata_hf() # Save updated metadata (removes deleted one)
        # For messages, we need to save the entire `all_chat_messages` dict again
        # This is inefficient but matches the original structure.
        # A better way would be to delete the specific chat file if each chat was a file.
        hf_token = os.getenv("HF_TOKEN")
        if hf_token:
            try:
                messages_filename_in_repo = f"user_memories/{self.user_id}_messages.json"
                messages_hf_path = f"datasets/{HF_USER_MEMORIES_DATASET_ID}/{messages_filename_in_repo}"
                with fs.open(messages_hf_path, "w", token=hf_token) as f: # This overwrites with current state
                    f.write(json.dumps(self.all_chat_messages, indent=2))
                print(f"Chat history for chat {chat_id} removed from HF by resaving all_chat_messages.")
            except Exception as e:
                print(f"Error resaving all_chat_messages after deletion: {e}")
                await cl.ErrorMessage(content=f"Error updating chat histories on cloud: {e}", author="System").send()

        print(f"Chat '{chat_id}' deleted.")

        if is_current_chat:
            if self.chat_metadata:
                await self.switch_chat(next(iter(self.chat_metadata)))
            else:
                await self.create_new_chat(is_initial_greeting=True)
        # UI should refresh based on action callbacks in ui.py

    async def rename_chat(self, chat_id: str, new_name: str):
        if not self.long_term_memory_enabled: return
        if chat_id and new_name and new_name != self.chat_metadata.get(chat_id):
            self.chat_metadata[chat_id] = new_name
            await self._save_chat_metadata_hf()
            print(f"Renamed chat '{chat_id}' to '{new_name}'")

    # --- Content Generation for Download ---
    def get_discussion_markdown(self, chat_id: str) -> str:
        messages_to_export = self.all_chat_messages.get(chat_id, [])
        # ... (same logic as original)
        markdown_content = []
        for msg in messages_to_export:
            role = msg["role"].capitalize()
            content = msg["content"]
            # Basic sanitization for markdown (e.g. if content itself has markdown)
            content = content.replace("*", "\\*").replace("_", "\\_") # Basic escape, might need more
            markdown_content.append(f"**{role}:**\n{content}\n\n---")
        return "\n".join(markdown_content)


    def get_discussion_docx(self, chat_id: str) -> bytes:
        messages_to_export = self.all_chat_messages.get(chat_id, [])
        document = Document()
        document.add_heading(f"Chat Discussion: {self.chat_metadata.get(chat_id, 'Untitled Chat')}", level=1)
        # ... (same logic as original)
        for msg in messages_to_export:
            role = msg["role"].capitalize()
            content = msg["content"]
            document.add_heading(f"{role}:", level=3)
            document.add_paragraph(content)
            document.add_paragraph("---")

        byte_stream = BytesIO()
        document.save(byte_stream)
        byte_stream.seek(0)
        return byte_stream.getvalue()

    # --- User Input Handling ---
    async def handle_user_input(self, user_query: str) -> AsyncGenerator[str, None]:
        """Processes user input and yields agent's response tokens."""
        if not self.current_chat_id and not self.chat_modified : # First message in a new session potentially
            # This check might be redundant if create_new_chat is always called first.
            # However, if a user sends a message when current_chat_id is None (e.g. after forget_me)
            await self.create_new_chat()
            # After creating a new chat, its metadata (name) should be saved if LTM is on.
            if self.long_term_memory_enabled and not self.chat_modified: # chat_modified would be false for a brand new chat
                 await self._save_chat_metadata_hf() # Save metadata of this new chat
                 self.chat_modified = True # Mark as modified so history saves next

        self.messages.append({"role": "user", "content": user_query})
        
        # Stream agent response
        full_response_content = ""
        async for token in self.get_agent_response_stream(user_query):
            full_response_content += token
            yield token # Yield token for ui.py to stream to cl.Message

        self.messages.append({"role": "assistant", "content": full_response_content})
        self.chat_modified = True # Mark that this chat has new messages

        if self.long_term_memory_enabled:
            await self._save_chat_history_hf() # Save full history of current chat

        self.suggested_prompts = await self.generate_suggested_prompts_async(self.messages)
        # The new suggested prompts will be available in self.suggested_prompts
        # ui.py can then decide to display them (e.g. after the assistant message)

    async def handle_regeneration_request(self) -> AsyncGenerator[str, None]:
        if not self.messages or self.messages[-1]['role'] != 'assistant':
            yield "Cannot regenerate: Last message not from assistant or no messages."
            return

        if len(self.messages) == 1: # Initial greeting
            # Regenerate greeting (it's stateless, so just call generator again)
            # This is a simplified regeneration; a true greeting regen might have specific logic
            self.initial_greeting_text = generate_llm_greeting() # Re-generate
            self.messages[0]['content'] = self.initial_greeting_text
            if self.long_term_memory_enabled: await self._save_chat_history_hf()
            self.suggested_prompts = await self.generate_suggested_prompts_async(self.messages)
            yield self.initial_greeting_text # Yield the whole new greeting as one "stream"
            return

        self.messages.pop() # Remove last assistant message
        if not self.messages or self.messages[-1]['role'] != 'user':
            yield "Cannot regenerate: No preceding user query found."
            # Restore popped message or add error to self.messages? For now, stream error.
            return

        last_user_query = self.messages[-1]['content']
        
        # History for regen is messages *before* the last user query
        # Effectively, self.messages now ends with the user query that led to the response we popped.
        # So, format_chat_history will use messages up to this user query.
        
        temp_last_user_msg = self.messages.pop() # Pop user message for history
        
        full_response_content = "" # Initialize the variable here
        async for token in self.get_agent_response_stream(last_user_query): # last_user_query is correct
            full_response_content += token
            yield token

        self.messages.append(temp_last_user_msg) # Add user message back
        self.messages.append({"role": "assistant", "content": full_response_content}) # Add new assistant response

        if self.long_term_memory_enabled: await self._save_chat_history_hf()
        self.suggested_prompts = await self.generate_suggested_prompts_async(self.messages)


    async def forget_me_and_reset(self):
        user_id_to_delete = self.user_id # Current user_id from AppState
        hf_token = os.getenv("HF_TOKEN")

        if user_id_to_delete and hf_token and self.long_term_memory_enabled:
            try:
                metadata_hf_path = f"datasets/{HF_USER_MEMORIES_DATASET_ID}/user_memories/{user_id_to_delete}_metadata.json"
                messages_hf_path = f"datasets/{HF_USER_MEMORIES_DATASET_ID}/user_memories/{user_id_to_delete}_messages.json"

                # Attempt to delete both files
                fs.rm([metadata_hf_path, messages_hf_path], token=hf_token, missing_ok=True) # missing_ok=True
                print(f"Attempted deletion of data for user '{user_id_to_delete}' from Hugging Face.")
            except Exception as e: # Catch other potential errors during delete
                print(f"Error deleting data for user {user_id_to_delete} from HF: {e}")
                await cl.ErrorMessage(content=f"Failed to delete data from cloud: {e}", author="System").send()

        # Reset AppState to a fresh state
        self.user_id = self.user_session.get("id") # Re-assign session ID, or handle new ID generation
        self.chat_metadata = {}
        self.all_chat_messages = {}
        self.current_chat_id = None
        self.messages = []
        self.suggested_prompts = DEFAULT_PROMPTS.copy()
        self.uploaded_documents = {}
        self.uploaded_dataframes = {}
        self.chat_modified = False

        # Also clear relevant items from cl.user_session to ensure a clean slate for Chainlit's UI
        self.user_session.set("current_chat_id", None)
        self.user_session.set("messages", []) # ui.py might use this key
        # If LTM preference was stored in user_session, user might want to reset that too, or not.
        # For now, LTM preference itself is not reset by "Forget Me".

        # Trigger a new chat creation for the UI
        await self.create_new_chat(is_initial_greeting=True)

        # Inform ui.py to potentially re-render or clear displays
        # This is implicitly handled by Chainlit's flow when actions complete or messages are sent.
        print(f"Session reset after forget_me. New chat started.")
        await cl.Message(content="All your data associated with the previous session has been cleared. A new session has started.", author="System").send()


    async def set_long_term_memory_preference(self, enabled: bool):
        if self.long_term_memory_enabled == enabled:
            return # No change

        self.long_term_memory_enabled = enabled
        self.user_session.set("long_term_memory_pref", enabled)
        print(f"Long-term memory preference set to: {enabled}")

        # If LTM is now enabled, load data.
        # If LTM is now disabled, clear persisted data from AppState (but not from HF).
        if enabled:
            if self.user_id: # Check user_id exists
                user_data = self._load_user_data_from_hf(self.user_id)
                self.chat_metadata = user_data.get("metadata", {})
                self.all_chat_messages = user_data.get("messages", {})
                if self.chat_metadata: # If history exists, switch to first/last known chat
                    last_chat_id = self.user_session.get("current_chat_id")
                    if last_chat_id and last_chat_id in self.chat_metadata:
                        self.current_chat_id = last_chat_id
                    else:
                        self.current_chat_id = next(iter(self.chat_metadata))
                    self.messages = self.all_chat_messages.get(self.current_chat_id, [])
                else: # No history, start new
                    await self.create_new_chat(is_initial_greeting=True)
            else: # No user_id, start new
                 await self.create_new_chat(is_initial_greeting=True)
        else: # LTM disabled
            # Clear local loaded history, start a new temporary chat
            self.chat_metadata = {}
            self.all_chat_messages = {}
            await self.create_new_chat(is_initial_greeting=True)

        self.suggested_prompts = await self.generate_suggested_prompts_async(self.messages)
        # ui.py will need to reflect this change, possibly by re-rendering elements.

    async def update_llm_settings(self, settings: Dict[str, Any]):
        if "llm_temperature" in settings: self.llm_temperature = float(settings["llm_temperature"])
        if "llm_verbosity" in settings: self.llm_verbosity = int(settings["llm_verbosity"])
        if "search_results_count" in settings: self.search_results_count = int(settings["search_results_count"])
        
        # Persist to user_session so ui.py can pick them up for display
        self.user_session.set("llm_temperature", self.llm_temperature)
        self.user_session.set("llm_verbosity", self.llm_verbosity)
        self.user_session.set("search_results_count", self.search_results_count)

        # Re-initialize agent if search_results_count changed, as it's a param to create_orchestrator_agent
        # This is a simplification; a more granular update would be better.
        if "search_results_count" in settings:
            await self._setup_agent_async()

        print(f"LLM settings updated: Temp={self.llm_temperature}, Verb={self.llm_verbosity}, SearchN={self.search_results_count}")
        await cl.Message(content="LLM settings updated.", author="System").send()

    async def generate_suggested_prompts_async(self, chat_history_list: List[Dict[str, Any]]) -> List[str]:
        # This needs to be async if the underlying generate_suggested_prompts is blocking
        # or uses network calls. For now, assume it's CPU-bound or fast enough.
        # If it's an LLM call, it should be async.
        # For this example, let's make it async.
        if _global_llm_initialization_error:
            return DEFAULT_PROMPTS.copy()

        prompts = await cl.make_async(generate_suggested_prompts)(chat_history_list)
        return prompts

    async def greet_user(self):
        """Sends the initial greeting message and sets up UI elements."""
        await cl.Message(content=self.initial_greeting_text, author="ESI").send()
        await ui.setup_actions_and_settings()

    async def display_suggested_prompts(self):
        """Displays suggested prompts as actions."""
        if self.suggested_prompts:
            # Added payload={} to the cl.Action constructor
            actions = [cl.Action(name=prompt, value=prompt, label=prompt, payload={}) for prompt in self.suggested_prompts]
            await cl.Message(content="Here are some things you can ask:", actions=actions).send()

# --- Chainlit Integration ---
# The main way to integrate is for ui.py's @cl.on_chat_start to create an AppState instance
# and store it in cl.user_session. Then, all other Chainlit callbacks in ui.py
# will retrieve this instance and call its methods.

@cl.on_chat_start
async def on_chat_start():
    """
    Initializes the application state and greets the user.
    This function is called when a new chat session starts.
    """
    # Create an instance of AppState, which also handles global LLM initialization
    app_state = AppState(cl.user_session)
    
    # Store the instance in the user session
    cl.user_session.set("app_state", app_state)
    
    # Now, call the async initialization method on the instance
    await app_state.initialize_session()
    
    # Set up the message history
    cl.user_session.set("messages", app_state.messages)
    
    # Greet the user and display prompts via the instance
    await app_state.greet_user()
    await app_state.display_suggested_prompts()
    
    print(f"app.py: @cl.on_chat_start: AppState initialized and stored in user_session.")


# The Chainlit CLI will be used to run the app, e.g., `chainlit run app.py -w`
# Chainlit will automatically discover @cl.on_chat_start, @cl.on_message from ui.py (due to import ui)
# and other decorators.

# Make sure environment variables (GOOGLE_API_KEY, HF_TOKEN, etc.) are loaded.
# load_dotenv() is already called at the top.

if not os.getenv("GOOGLE_API_KEY"):
    print("⚠️ WARNING: GOOGLE_API_KEY environment variable not set. The agent may not work properly.")
if not os.getenv("HF_TOKEN"):
    print("⚠️ WARNING: HF_TOKEN environment variable not set. Long-term memory features may not work.")

# The rest of the Streamlit-specific main() function and its setup is removed.
# Chainlit manages the application lifecycle and event loop.
